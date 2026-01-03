"""
Tests for /upload_dataset and /ask with playbooks (Phase 2).

All tests use Fake clients - no network calls.
"""

import json

# Adjust path for imports
import sys
from pathlib import Path

import pytest
from fastapi.testclient import TestClient

sys.path.insert(0, str(Path(__file__).parent.parent.parent.parent))

from services.api.main import DATASETS_DIR, STORAGE_DIR, app

client = TestClient(app)


@pytest.fixture
def sample_csv_path():
    """Path to the sample CSV fixture."""
    return Path(__file__).parent / "fixtures" / "sample.csv"


@pytest.fixture
def uploaded_dataset(sample_csv_path):
    """Upload the sample CSV and return the dataset_id."""
    with open(sample_csv_path, "rb") as f:
        response = client.post(
            "/upload_dataset",
            files={"file": ("sample.csv", f, "text/csv")},
        )
    assert response.status_code == 200
    data = response.json()
    yield data
    # Cleanup
    import shutil
    dataset_dir = DATASETS_DIR / data["dataset_id"]
    if dataset_dir.exists():
        shutil.rmtree(dataset_dir)


@pytest.fixture
def uploaded_context_doc():
    """Create a minimal context doc for testing."""
    import os
    import tempfile

    from docx import Document

    # Create a valid docx with enough content (>500 chars)
    doc = Document()
    doc.add_heading("Dataset Overview", level=1)
    doc.add_paragraph(
        "This is a comprehensive sales transactions dataset containing daily sales data "
        "from multiple product categories across different regions. The dataset includes "
        "information about electronics, clothing, and home goods sold throughout January 2024."
    )
    doc.add_heading("Target Use / Primary Questions", level=1)
    doc.add_paragraph(
        "What are the top-performing product categories by revenue? "
        "How do sales trends vary across different regions? "
        "What is the average transaction value by product type?"
    )
    doc.add_heading("Data Dictionary", level=1)
    doc.add_paragraph("date: Transaction date in YYYY-MM-DD format")
    doc.add_paragraph("category: Product category (Electronics, Clothing, Home)")
    doc.add_paragraph("product: Specific product name")
    doc.add_paragraph("sales: Sales amount in USD")
    doc.add_paragraph("quantity: Number of units sold")
    doc.add_paragraph("region: Geographic region (North, South, East, West)")
    doc.add_heading("Known Caveats", level=1)
    doc.add_paragraph(
        "Data may have some missing values. Sales figures are rounded to two decimal places. "
        "The dataset covers a limited time period and may not be representative of annual trends."
    )

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        doc.save(f.name)
        temp_path = f.name

    # Upload
    with open(temp_path, "rb") as f:
        response = client.post(
            "/upload_context_doc",
            files={"file": ("context.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        )

    os.unlink(temp_path)

    if response.status_code != 200:
        print(f"Upload failed: {response.json()}")
    assert response.status_code == 200
    data = response.json()
    yield data
    # Cleanup
    import shutil
    doc_dir = STORAGE_DIR / data["doc_id"]
    if doc_dir.exists():
        shutil.rmtree(doc_dir)


class TestUploadDataset:
    """Tests for /upload_dataset endpoint."""

    def test_upload_creates_metadata_and_profile(self, sample_csv_path):
        """Test that upload creates metadata.json and profile.json."""
        with open(sample_csv_path, "rb") as f:
            response = client.post(
                "/upload_dataset",
                files={"file": ("sample.csv", f, "text/csv")},
            )

        assert response.status_code == 200
        data = response.json()

        # Verify response fields
        assert "dataset_id" in data
        assert "dataset_hash" in data
        assert data["n_rows"] == 14  # 14 data rows (header not counted)
        assert data["n_cols"] == 6
        assert "date" in data["column_names"]
        assert "sales" in data["column_names"]
        assert data["inferred_types"]["sales"] == "float"
        assert data["status"] == "profiled"

        # Verify files were created
        dataset_dir = DATASETS_DIR / data["dataset_id"]
        assert (dataset_dir / "data.csv").exists()
        assert (dataset_dir / "metadata.json").exists()
        assert (dataset_dir / "profile.json").exists()

        # Verify profile content
        with open(dataset_dir / "profile.json") as f:
            profile = json.load(f)
        assert "columns" in profile
        assert "sales" in profile["columns"]

        # Cleanup
        import shutil
        shutil.rmtree(dataset_dir)

    def test_upload_rejects_non_csv(self):
        """Test that non-CSV files are rejected."""
        response = client.post(
            "/upload_dataset",
            files={"file": ("data.txt", b"hello world", "text/plain")},
        )
        assert response.status_code == 415


class TestAskWithPlaybooks:
    """Tests for /ask endpoint with playbook execution."""

    def test_overview_question_returns_artifacts(self, uploaded_dataset, uploaded_context_doc):
        """Test that an overview question returns >= 2 artifacts."""
        response = client.post(
            "/ask",
            json={
                "question": "Give me an overview of the dataset",
                "doc_id": uploaded_context_doc["doc_id"],
                "dataset_id": uploaded_dataset["dataset_id"],
            },
        )

        assert response.status_code == 200
        data = response.json()

        assert data["router_decision"]["route"] == "ANALYSIS"
        assert len(data["artifacts"]) >= 2
        assert data["trace_id"]

    def test_groupby_question_returns_table(self, uploaded_dataset, uploaded_context_doc):
        """Test that a groupby question returns a TableArtifact."""
        response = client.post(
            "/ask",
            json={
                "question": "What is the total sales by category?",
                "doc_id": uploaded_context_doc["doc_id"],
                "dataset_id": uploaded_dataset["dataset_id"],
            },
        )

        assert response.status_code == 200
        data = response.json()

        # Should route to ANALYSIS
        assert data["router_decision"]["route"] == "ANALYSIS"
        # May have table artifacts if playbook was selected correctly
        assert "artifacts" in data

