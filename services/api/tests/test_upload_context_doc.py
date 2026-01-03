"""
Tests for /upload_context_doc endpoint.
"""

import io
import sys
from pathlib import Path

from docx import Document
from fastapi.testclient import TestClient

sys.path.insert(0, str(Path(__file__).parent.parent.parent.parent))

from services.api.main import app

client = TestClient(app)


def create_docx_bytes(content_dict: dict) -> bytes:
    """Helper to create a .docx file in memory."""
    doc = Document()

    for heading, paragraphs in content_dict.items():
        doc.add_heading(heading, level=1)
        if isinstance(paragraphs, list):
            for para in paragraphs:
                doc.add_paragraph(para)
        else:
            doc.add_paragraph(paragraphs)

    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


def test_rejects_non_docx():
    """Test that non-.docx files are rejected."""
    # Create a fake text file
    fake_file = io.BytesIO(b"This is not a docx file")

    response = client.post(
        "/upload_context_doc",
        files={"file": ("test.txt", fake_file, "text/plain")}
    )

    assert response.status_code == 415
    assert "docx" in response.json()["detail"]["error"].lower()


def test_rejects_missing_headings():
    """Test that documents missing required headings are rejected."""
    # Create a docx with only some headings
    content = {
        "Dataset Overview": "This is an overview of the dataset.",
        "Target Use / Primary Questions": "What are the main questions?",
        # Missing: "Data Dictionary" and "Known Caveats"
    }

    docx_bytes = create_docx_bytes(content)

    response = client.post(
        "/upload_context_doc",
        files={"file": ("test.docx", io.BytesIO(docx_bytes), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
    )

    assert response.status_code == 422
    data = response.json()
    assert "missing" in data["detail"]["error"].lower()
    assert "Data Dictionary" in data["detail"]["missing_headings"]
    assert "Known Caveats" in data["detail"]["missing_headings"]


def test_rejects_too_short_document():
    """Test that documents with insufficient content are rejected."""
    # Create a docx with all headings but minimal content
    content = {
        "Dataset Overview": "Short.",
        "Target Use / Primary Questions": "Q?",
        "Data Dictionary": "D.",
        "Known Caveats": "C.",
    }

    docx_bytes = create_docx_bytes(content)

    response = client.post(
        "/upload_context_doc",
        files={"file": ("test.docx", io.BytesIO(docx_bytes), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
    )

    assert response.status_code == 422
    data = response.json()
    assert "too short" in data["detail"]["error"].lower() or "minimum" in data["detail"]["error"].lower()


def test_accepts_valid_document():
    """Test that a valid document is accepted and processed correctly."""
    # Create a valid docx with all headings and sufficient content
    content = {
        "Dataset Overview": (
            "This dataset contains customer transaction records from our e-commerce platform "
            "for Q1 2024. Each row represents a single transaction. The data includes 50,000 "
            "transactions from 12,000 unique customers across 5 product categories. The dataset "
            "was extracted from our production PostgreSQL database on April 1, 2024."
        ),
        "Target Use / Primary Questions": (
            "This dataset is intended to support quarterly business review analysis. "
            "Primary questions include: What are the top-performing product categories? "
            "Which customer segments have the highest lifetime value? What factors correlate "
            "with purchase frequency? Are there seasonal patterns in transaction volume?"
        ),
        "Data Dictionary": (
            "transaction_id (string): Unique transaction identifier. "
            "customer_id (string): Unique customer identifier. "
            "transaction_date (date): Date of transaction. "
            "product_category (string): Category of purchased product. "
            "amount_usd (decimal): Transaction amount in USD. "
            "payment_method (string): Payment method used. "
            "is_first_purchase (boolean): Whether this is customer's first purchase. "
            "discount_applied (decimal): Discount amount in USD. "
            "shipping_country (string): Country code for shipping. "
            "order_status (string): Current order status."
        ),
        "Known Caveats": (
            "Approximately 5% of transactions have missing customer_id values. "
            "This dataset only includes completed transactions. Refunds are recorded "
            "as separate transactions with negative amounts. International transactions "
            "are converted to USD using daily exchange rates."
        ),
    }

    docx_bytes = create_docx_bytes(content)

    response = client.post(
        "/upload_context_doc",
        files={"file": ("test.docx", io.BytesIO(docx_bytes), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
    )

    assert response.status_code == 200
    data = response.json()

    # Verify response structure
    assert data["status"] == "indexed"
    assert "doc_id" in data
    assert "doc_hash" in data
    assert data["num_chars"] >= 800
    assert data["num_chunks"] > 0
    assert data["errors"] is None


def test_health_endpoint():
    """Test that health endpoint works."""
    response = client.get("/health")
    assert response.status_code == 200
    assert response.json() == {"status": "ok"}


def test_version_endpoint():
    """Test that version endpoint returns expected fields."""
    response = client.get("/version")
    assert response.status_code == 200

    data = response.json()
    # Verify required fields are present
    assert "git_sha" in data
    assert "build_time" in data
    assert "app_env" in data

    # Verify types
    assert isinstance(data["git_sha"], str)
    assert isinstance(data["build_time"], str)
    assert isinstance(data["app_env"], str)

    # In test environment, app_env should be dev
    assert data["app_env"] == "dev"
