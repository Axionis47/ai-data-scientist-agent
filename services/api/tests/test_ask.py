"""
Tests for /ask endpoint with RAG agent.
Uses fake clients - no network calls.
"""

import io
import sys
from pathlib import Path

from docx import Document
from fastapi.testclient import TestClient

sys.path.insert(0, str(Path(__file__).parent.parent.parent.parent))

from services.api.main import app

client = TestClient(app)


def create_valid_docx_bytes() -> bytes:
    """Create a valid .docx file for testing."""
    doc = Document()

    content = {
        "Dataset Overview": (
            "This dataset contains customer transaction records from our e-commerce platform "
            "for Q1 2024. Each row represents a single transaction. The data includes 50,000 "
            "transactions from 12,000 unique customers across 5 product categories. "
            "The unit of observation is a single completed purchase transaction."
        ),
        "Target Use / Primary Questions": (
            "This dataset is intended to support quarterly business review analysis. "
            "Primary questions include: What are the top-performing product categories? "
            "Which customer segments have the highest lifetime value? What factors correlate "
            "with purchase frequency? Are there seasonal patterns in transaction volume?"
        ),
        "Data Dictionary": (
            "transaction_id (string): Unique transaction identifier. "
            "customer_id (string): Unique customer identifier. "
            "transaction_date (date): Date of transaction. "
            "product_category (string): Category of purchased product. "
            "amount_usd (decimal): Transaction amount in USD. "
            "payment_method (string): Payment method used. "
            "is_first_purchase (boolean): Whether this is customer's first purchase. "
            "discount_applied (decimal): Discount amount in USD. "
            "shipping_country (string): Country code for shipping. "
            "order_status (string): Current order status."
        ),
        "Known Caveats": (
            "Approximately 5% of transactions have missing customer_id values. "
            "This dataset only includes completed transactions. Refunds are recorded "
            "as separate transactions with negative amounts. International transactions "
            "are converted to USD using daily exchange rates."
        ),
    }

    for heading, paragraph in content.items():
        doc.add_heading(heading, level=1)
        doc.add_paragraph(paragraph)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


def upload_test_doc() -> str:
    """Upload a test document and return doc_id."""
    docx_bytes = create_valid_docx_bytes()
    response = client.post(
        "/upload_context_doc",
        files={"file": ("test.docx", io.BytesIO(docx_bytes), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
    )
    assert response.status_code == 200
    return response.json()["doc_id"]


def test_ask_returns_valid_response():
    """Test that /ask returns a schema-valid AskQuestionResponse."""
    doc_id = upload_test_doc()

    response = client.post(
        "/ask",
        json={
            "question": "What product categories are in the dataset?",
            "doc_id": doc_id,
        }
    )

    assert response.status_code == 200
    data = response.json()

    # Validate response structure
    assert "answer_text" in data
    assert "router_decision" in data
    assert "artifacts" in data
    assert "trace_id" in data

    # Validate router_decision
    assert data["router_decision"]["route"] in ["ANALYSIS", "CAUSAL", "REPORTING", "SYSTEM", "NEEDS_CLARIFICATION"]
    assert isinstance(data["router_decision"]["confidence"], float)
    assert isinstance(data["router_decision"]["reasons"], list)

    # Answer should be non-empty
    assert len(data["answer_text"]) > 0


def test_ask_includes_retrieved_chunks():
    """Test that /ask includes retrieved chunk artifact for matching question."""
    doc_id = upload_test_doc()

    response = client.post(
        "/ask",
        json={
            "question": "What is the unit of observation in this dataset?",
            "doc_id": doc_id,
        }
    )

    assert response.status_code == 200
    data = response.json()

    # Should have artifacts with retrieved chunks
    assert len(data["artifacts"]) > 0

    # At least one artifact should be text type with chunk info
    text_artifacts = [a for a in data["artifacts"] if a.get("type") == "text"]
    assert len(text_artifacts) > 0
    assert "chunk" in text_artifacts[0]["content"].lower()


def test_ask_causal_question_routes_to_needs_clarification():
    """Test that causal questions return NEEDS_CLARIFICATION with checklist."""
    doc_id = upload_test_doc()

    response = client.post(
        "/ask",
        json={
            "question": "What is the effect of discounts on purchase frequency?",
            "doc_id": doc_id,
        }
    )

    assert response.status_code == 200
    data = response.json()

    # Should route to NEEDS_CLARIFICATION for causal questions in Phase 1
    assert data["router_decision"]["route"] == "NEEDS_CLARIFICATION"

    # Should include a checklist artifact
    checklist_artifacts = [a for a in data["artifacts"] if a.get("type") == "checklist"]
    assert len(checklist_artifacts) > 0
    assert len(checklist_artifacts[0]["items"]) > 0


def test_ask_missing_doc_returns_404():
    """Test that /ask returns 404 for non-existent doc_id."""
    response = client.post(
        "/ask",
        json={
            "question": "What is in the dataset?",
            "doc_id": "non-existent-doc-id",
        }
    )

    assert response.status_code == 404

