"""
Script to create a valid sample context document in .docx format.
"""

from pathlib import Path

from docx import Document

# Create document
doc = Document()

# Add title
doc.add_heading("Sample Dataset Context Document", 0)

# Dataset Overview
doc.add_heading("Dataset Overview", 1)
doc.add_paragraph(
    "This dataset contains customer transaction records from our e-commerce platform "
    "for Q1 2024. Each row represents a single transaction. The data includes 50,000 "
    "transactions from 12,000 unique customers across 5 product categories. The dataset "
    "was extracted from our production PostgreSQL database on April 1, 2024. All personally "
    "identifiable information has been anonymized using SHA-256 hashing. The unit of observation "
    "is a single transaction, with each transaction potentially containing multiple line items "
    "aggregated into a single total amount."
)

# Target Use / Primary Questions
doc.add_heading("Target Use / Primary Questions", 1)
doc.add_paragraph(
    "This dataset is intended to support quarterly business review analysis and strategic "
    "planning for the product and marketing teams. Primary questions include:"
)
doc.add_paragraph("• What are the top-performing product categories by revenue and volume?", style="List Bullet")
doc.add_paragraph("• Which customer segments have the highest lifetime value and retention rates?", style="List Bullet")
doc.add_paragraph("• What factors correlate with purchase frequency and basket size?", style="List Bullet")
doc.add_paragraph("• Are there seasonal or weekly patterns in transaction volume?", style="List Bullet")
doc.add_paragraph("• How effective are discount campaigns at driving incremental revenue?", style="List Bullet")
doc.add_paragraph("• What is the geographic distribution of our customer base?", style="List Bullet")

# Data Dictionary
doc.add_heading("Data Dictionary", 1)
doc.add_paragraph("The dataset contains the following columns:")

columns = [
    ("transaction_id", "string", "Unique transaction identifier", "Primary key, format: TXN-XXXXXXXX"),
    ("customer_id", "string", "Unique customer identifier", "Foreign key, anonymized hash"),
    ("transaction_date", "date", "Date of transaction", "Format: YYYY-MM-DD, range: 2024-01-01 to 2024-03-31"),
    ("product_category", "string", "Category of purchased product", "Values: Electronics, Clothing, Home, Books, Sports"),
    ("amount_usd", "decimal", "Transaction amount in USD", "Always positive, range: $5.00 to $2,500.00"),
    ("payment_method", "string", "Payment method used", "Values: credit_card, debit_card, paypal, apple_pay"),
    ("is_first_purchase", "boolean", "Whether this is customer\'s first purchase", "Values: true, false"),
    ("discount_applied", "decimal", "Discount amount in USD", "0 if no discount, max 50% of amount_usd"),
    ("shipping_country", "string", "Country code for shipping", "ISO 2-letter code, primarily US, CA, UK, DE"),
    ("order_status", "string", "Current order status", "Values: pending, shipped, delivered, cancelled"),
    ("customer_segment", "string", "Customer segment classification", "Values: new, regular, vip, at_risk"),
    ("referral_source", "string", "How customer found us", "Values: organic, paid_search, social, email, direct"),
]

for col_name, col_type, description, notes in columns:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(f"{col_name}").bold = True
    p.add_run(f" ({col_type}): {description}. {notes}")

# Known Caveats
doc.add_heading("Known Caveats", 1)
doc.add_paragraph(
    "Users should be aware of the following limitations and data quality considerations:"
)
doc.add_paragraph(
    "• Approximately 5% of transactions have missing customer_id values, representing guest "
    "checkouts where customers did not create an account.",
    style="List Bullet"
)
doc.add_paragraph(
    "• This dataset only includes completed transactions. Abandoned carts, browsing sessions, "
    "and incomplete checkouts are not represented.",
    style="List Bullet"
)
doc.add_paragraph(
    "• Refunds and returns are recorded as separate transactions with negative amounts, not "
    "as modifications to the original transaction record.",
    style="List Bullet"
)
doc.add_paragraph(
    "• International transactions are converted to USD using the daily exchange rate at the "
    "time of transaction. Exchange rate fluctuations may affect period-over-period comparisons.",
    style="List Bullet"
)
doc.add_paragraph(
    "• The product category taxonomy was updated on March 1, 2024. Earlier transactions have "
    "been retroactively mapped to the new category structure, which may introduce minor "
    "classification inconsistencies.",
    style="List Bullet"
)
doc.add_paragraph(
    "• Customer segment classifications are computed monthly and reflect the segment at the "
    "time of transaction, not the current segment.",
    style="List Bullet"
)

# Save document
output_dir = Path(__file__).parent.parent / "docs" / "samples"
output_dir.mkdir(parents=True, exist_ok=True)
output_path = output_dir / "context_template.docx"

doc.save(output_path)
print(f"Created sample document at: {output_path}")
print(f"File size: {output_path.stat().st_size} bytes")

